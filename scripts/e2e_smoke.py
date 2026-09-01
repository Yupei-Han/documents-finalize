#!/usr/bin/env python3
"""Run an isolated positive/negative end-to-end regression for both document skills."""

from __future__ import annotations

import argparse
import hashlib
import json
import os
import shutil
import subprocess
import sys
import tempfile
from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
FAST_ROOT = ROOT.parent / "documents-fast"


def sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def run(label: str, command: list[str], *, expect_success: bool = True) -> subprocess.CompletedProcess[str]:
    result = subprocess.run(
        command,
        capture_output=True,
        text=True,
        encoding="utf-8",
        errors="replace",
        check=False,
    )
    passed = result.returncode == 0 if expect_success else result.returncode != 0
    if not passed:
        raise RuntimeError(
            f"{label} returned {result.returncode}; expected "
            f"{'success' if expect_success else 'failure'}\n"
            f"STDOUT:\n{result.stdout}\nSTDERR:\n{result.stderr}"
        )
    print(f"[PASS] {label}")
    return result


def create_fixture(path: Path) -> None:
    document = Document()
    document.add_heading("Documents split regression", level=1)
    document.add_paragraph(
        "This deterministic working fixture verifies source preservation, evidence binding, "
        "page coverage, and fail-closed release behavior."
    )
    table = document.add_table(rows=2, cols=2)
    table.style = "Table Grid"
    table.cell(0, 0).text = "Check"
    table.cell(0, 1).text = "Expected"
    table.cell(1, 0).text = "Status"
    table.cell(1, 1).text = "PASS"
    document.save(path)


def python(script: Path, *args: str) -> list[str]:
    return [sys.executable, str(script), *args]


def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--renderer", choices=["libreoffice", "word"], default="libreoffice")
    parser.add_argument("--node", type=Path)
    parser.add_argument("--pwsh", type=Path)
    parser.add_argument("--pdftoppm", type=Path)
    parser.add_argument(
        "--workdir",
        type=Path,
        help="Optional new directory to retain; default uses a deleted temporary directory",
    )
    args = parser.parse_args()
    if not FAST_ROOT.is_dir():
        raise SystemExit(f"peer documents-fast skill not found: {FAST_ROOT}")
    node = args.node or Path(shutil.which("node") or "")
    if not node.is_file():
        raise SystemExit("smoke test requires a valid --node path")

    temp_owner: tempfile.TemporaryDirectory[str] | None = None
    if args.workdir:
        root = args.workdir.resolve()
        if root.exists():
            raise SystemExit(f"--workdir must be new: {root}")
        root.mkdir(parents=True)
    else:
        temp_owner = tempfile.TemporaryDirectory(prefix="documents_split_e2e_")
        root = Path(temp_owner.name)

    try:
        working = root / "working"
        qa = root / "qa"
        delivery = root / "delivery"
        working.mkdir()
        qa.mkdir()
        delivery.mkdir()
        source = root / "source.docx"
        candidate = working / "candidate.docx"
        final = delivery / "release.docx"
        create_fixture(source)
        fast_start = qa / "candidate-operation-start.json"
        fast_completion = qa / "candidate-operation-complete.json"
        final_start = qa / "release-operation-start.json"
        final_completion = qa / "release-operation-complete.json"
        run(
            "fast operation start receipt",
            [
                str(node),
                str(FAST_ROOT / "container_tools" / "mark_artifact_operation_started.mjs"),
                "create",
                "--operation-kind",
                "edit",
                "--expected-output-count",
                "1",
                "--output-format",
                "docx",
                "--receipt",
                str(fast_start),
                "--input",
                str(source),
                "--planned-output",
                str(candidate),
            ],
        )
        shutil.copyfile(source, candidate)
        run(
            "fast operation completion receipt",
            [
                str(node),
                str(FAST_ROOT / "container_tools" / "mark_artifact_operation_started.mjs"),
                "verify",
                "--receipt",
                str(fast_start),
                "--require-outputs",
                "--verification",
                str(fast_completion),
            ],
        )
        run(
            "duplicate operation receipt rejection",
            [
                str(node),
                str(FAST_ROOT / "container_tools" / "mark_artifact_operation_started.mjs"),
                "create",
                "--operation-kind",
                "edit",
                "--expected-output-count",
                "1",
                "--output-format",
                "docx",
                "--receipt",
                str(fast_start),
                "--input",
                str(source),
                "--planned-output",
                str(candidate),
            ],
            expect_success=False,
        )
        run(
            "final operation start receipt",
            [
                str(node),
                str(ROOT / "container_tools" / "mark_artifact_operation_started.mjs"),
                "create",
                "--operation-kind",
                "edit",
                "--expected-output-count",
                "1",
                "--output-format",
                "docx",
                "--receipt",
                str(final_start),
                "--input",
                str(candidate),
                "--planned-output",
                str(final),
            ],
        )
        shutil.copyfile(candidate, final)
        run(
            "final operation completion receipt",
            [
                str(node),
                str(ROOT / "container_tools" / "mark_artifact_operation_started.mjs"),
                "verify",
                "--receipt",
                str(final_start),
                "--require-outputs",
                "--verification",
                str(final_completion),
            ],
        )
        source_hash = sha256(source)

        fast_scripts = FAST_ROOT / "scripts"
        final_scripts = ROOT / "scripts"
        source_risk = qa / "source-risk-fast.json"
        candidate_risk = qa / "candidate-risk-fast.json"
        handoff = qa / "working.handoff.json"
        run(
            "fast source risk scan",
            python(fast_scripts / "docx_risk_scan.py", str(source), "--output", str(source_risk)),
        )
        run(
            "fast candidate risk scan",
            python(
                fast_scripts / "docx_risk_scan.py",
                str(candidate),
                "--output",
                str(candidate_risk),
            ),
        )
        run(
            "fast handoff creation",
            python(
                fast_scripts / "document_handoff.py",
                "create",
                "--source",
                str(source),
                "--source-risk-report",
                str(source_risk),
                "--candidate",
                str(candidate),
                "--candidate-risk-report",
                str(candidate_risk),
                "--output",
                str(handoff),
                "--authorized-mode",
                "faithful editing",
                "--summary",
                "regression fixture",
                "--method",
                "exact-byte copy",
                "--checked",
                "package and risk report",
                "--pending",
                "documents-finalize release gates",
            ),
        )
        run(
            "fast handoff live verification",
            python(
                final_scripts / "document_handoff.py",
                "verify",
                str(handoff),
                "--source",
                str(source),
                "--candidate",
                str(candidate),
            ),
        )
        run(
            "fast overwrite rejection",
            python(fast_scripts / "privacy_scrub.py", str(source), "--out", str(candidate)),
            expect_success=False,
        )

        preflight_dir = qa / "preflight"
        preflight = preflight_dir / "preflight.json"
        run(
            "final preflight",
            python(
                final_scripts / "document_preflight.py",
                "--source",
                str(source),
                "--candidate",
                str(candidate),
                "--final",
                str(final),
                "--output-dir",
                str(preflight_dir),
            ),
        )

        render_dir = qa / "render"
        manifest = render_dir / "renderer.json"
        if args.renderer == "libreoffice":
            render_command = python(
                final_scripts / "render_docx_libreoffice.py",
                "--input",
                str(final),
                "--output-dir",
                str(render_dir),
                "--manifest",
                str(manifest),
            )
        else:
            pwsh = args.pwsh or Path(
                r"C:\Program Files\PowerShell\7\pwsh.exe"
                if os.name == "nt"
                else shutil.which("pwsh") or "pwsh"
            )
            pdftoppm = args.pdftoppm or Path(shutil.which("pdftoppm") or "")
            if not pwsh.is_file() or not pdftoppm.is_file():
                raise RuntimeError("Word smoke test requires valid --pwsh and --pdftoppm paths")
            render_command = [
                str(pwsh),
                "-NoProfile",
                "-File",
                str(final_scripts / "render_docx_word.ps1"),
                "-InputPath",
                str(final),
                "-OutputDir",
                str(render_dir),
                "-ManifestPath",
                str(manifest),
                "-PdfToPpmPath",
                str(pdftoppm),
            ]
        run(f"{args.renderer} manifest render", render_command)
        manifest_payload = json.loads(manifest.read_text(encoding="utf-8"))
        repair_observation = manifest_payload["application_open"]["repair_observation"]
        if (
            repair_observation.get("status") != "NO_DIAGNOSTIC_OBSERVED"
            or repair_observation.get("observed_warning_count") != 0
            or repair_observation.get("absence_proven") is not False
            or not str(repair_observation.get("capture_scope", "")).strip()
        ):
            raise RuntimeError("renderer repair observation semantics are invalid")
        print("[PASS] renderer repair observation semantics")
        pages = [entry["page"] for entry in manifest_payload["pages"]]
        if not pages:
            raise RuntimeError("renderer manifest contains no pages")

        page_review = qa / "page-review.json"
        missing_page = qa / "page-review-missing.json"
        run(
            "missing-page rejection",
            python(
                final_scripts / "document_release.py",
                "page-review",
                "--document",
                str(final),
                "--renderer-manifest",
                str(manifest),
                "--output",
                str(missing_page),
                "--page-result",
                "999=PASS: deliberately wrong page",
                "--method",
                "negative regression",
            ),
            expect_success=False,
        )
        page_command = python(
            final_scripts / "document_release.py",
            "page-review",
            "--document",
            str(final),
            "--renderer-manifest",
            str(manifest),
            "--output",
            str(page_review),
            "--method",
            "full-size image inspection regression",
        )
        for page in pages:
            page_command.extend(
                ["--page-result", f"{page}=PASS: full page decoded and inspected"]
            )
        run("complete per-page review", page_command)

        render_evidence = qa / "render-evidence.json"
        run(
            "render evidence",
            python(
                final_scripts / "document_release.py",
                "render-evidence",
                "--document",
                str(final),
                "--renderer-manifest",
                str(manifest),
                "--page-review",
                str(page_review),
                "--output",
                str(render_evidence),
            ),
        )

        scope = qa / "scope.json"
        structure = qa / "structure.json"
        native = qa / "native.json"
        application = qa / "application.json"
        run(
            "scope review",
            python(
                final_scripts / "document_release.py",
                "scope-review",
                "--preflight",
                str(preflight),
                "--output",
                str(scope),
                "--authorized-mode",
                "faithful editing",
                "--method",
                "source/candidate/final exact-byte and requirement review",
                "--requirement-result",
                "fixture=PASS: candidate and final match the authorized fixture",
            ),
        )
        run(
            "structure review",
            python(
                final_scripts / "document_release.py",
                "structure-review",
                "--preflight",
                str(preflight),
                "--output",
                str(structure),
                "--method",
                "live package-part and risk-delta coverage",
            ),
        )
        run(
            "native policy review",
            python(
                final_scripts / "document_release.py",
                "native-policy-review",
                "--preflight",
                str(preflight),
                "--output",
                str(native),
                "--comments-policy",
                "preserve",
                "--revisions-policy",
                "preserve",
                "--fields-policy",
                "preserve-editable",
                "--citations-policy",
                "none",
                "--method",
                "fresh candidate/final risk comparison",
            ),
        )
        run(
            "application review",
            python(
                final_scripts / "document_release.py",
                "application-review",
                "--preflight",
                str(preflight),
                "--renderer-manifest",
                str(manifest),
                "--output",
                str(application),
                "--method",
                "manifest-derived open/export evidence",
            ),
        )

        delta_dir = root / "delta"
        delta_dir.mkdir()
        delta_source = delta_dir / "source.docx"
        delta_candidate = delta_dir / "candidate.docx"
        delta_final = delta_dir / "final.docx"
        shutil.copyfile(source, delta_source)
        shutil.copyfile(source, delta_candidate)
        changed_document = Document(str(delta_candidate))
        changed_document.add_paragraph("Authorized visible-content delta for coverage testing.")
        changed_document.save(delta_final)
        run(
            "release-only candidate/final mismatch rejection",
            python(
                final_scripts / "document_preflight.py",
                "--source",
                str(delta_source),
                "--candidate",
                str(delta_candidate),
                "--final",
                str(delta_final),
                "--output-dir",
                str(qa / "release-only-rejection"),
            ),
            expect_success=False,
        )
        shutil.copyfile(delta_final, delta_candidate)
        delta_preflight_dir = qa / "delta-preflight"
        delta_preflight = delta_preflight_dir / "preflight.json"
        run(
            "delta preflight",
            python(
                final_scripts / "document_preflight.py",
                "--source",
                str(delta_source),
                "--candidate",
                str(delta_candidate),
                "--final",
                str(delta_final),
                "--output-dir",
                str(delta_preflight_dir),
            ),
        )
        run(
            "visible-text review omission rejection",
            python(
                final_scripts / "document_release.py",
                "scope-review",
                "--preflight",
                str(delta_preflight),
                "--output",
                str(qa / "delta-scope-missing.json"),
                "--authorized-mode",
                "faithful editing",
                "--method",
                "negative coverage regression",
                "--requirement-result",
                "delta=PASS: requested text addition",
            ),
            expect_success=False,
        )
        run(
            "visible-text review coverage",
            python(
                final_scripts / "document_release.py",
                "scope-review",
                "--preflight",
                str(delta_preflight),
                "--output",
                str(qa / "delta-scope.json"),
                "--authorized-mode",
                "faithful editing",
                "--method",
                "explicit semantic delta review",
                "--requirement-result",
                "delta=PASS: requested text addition",
                "--text-change-review",
                "visible_content=PASS: added paragraph matches the authorized test",
            ),
        )
        run(
            "changed-part coverage omission rejection",
            python(
                final_scripts / "document_release.py",
                "structure-review",
                "--preflight",
                str(delta_preflight),
                "--output",
                str(qa / "delta-structure-missing.json"),
                "--method",
                "negative coverage regression",
            ),
            expect_success=False,
        )
        comparison_payload = json.loads(
            (delta_preflight_dir / "comparison.json").read_text(encoding="utf-8")
        )
        package_delta = comparison_payload["comparisons"]["baseline_to_final"]["package"]
        changed_parts = sorted(
            set(package_delta["changed_parts"])
            | set(package_delta["added_parts"])
            | set(package_delta["removed_parts"])
        )
        delta_preflight_payload = json.loads(delta_preflight.read_text(encoding="utf-8"))
        risk_delta_keys: list[str] = []
        for group in ("counts", "feature_flags"):
            left = delta_preflight_payload["risk"]["source"][group]
            right = delta_preflight_payload["risk"]["final"][group]
            risk_delta_keys.extend(
                f"{group}.{key}"
                for key in sorted(set(left) | set(right))
                if left.get(key) != right.get(key)
            )
        structure_command = python(
            final_scripts / "document_release.py",
            "structure-review",
            "--preflight",
            str(delta_preflight),
            "--output",
            str(qa / "delta-structure.json"),
            "--method",
            "complete changed-part and risk-delta review",
        )
        for part in changed_parts:
            structure_command.extend(["--reviewed-part", part])
        for key in risk_delta_keys:
            structure_command.extend(
                ["--delta-explanation", f"{key}=PASS: authorized regression delta"]
            )
        run("complete changed-part coverage", structure_command)

        if args.renderer == "libreoffice":
            tracked_fixture_dir = root / "tracked-fixtures"
            run(
                "tracked fixture generation",
                python(
                    fast_scripts / "make_fixtures.py",
                    "--outdir",
                    str(tracked_fixture_dir),
                    "--only",
                    "tracked",
                ),
            )
            tracked_fixture = tracked_fixture_dir / "tracked_changes_fixture.docx"
            tracked_source = root / "tracked-source.docx"
            tracked_candidate = working / "tracked-candidate.docx"
            tracked_final = working / "tracked-final.docx"
            shutil.copyfile(tracked_fixture, tracked_source)
            shutil.copyfile(tracked_fixture, tracked_candidate)
            shutil.copyfile(tracked_fixture, tracked_final)
            tracked_preflight_dir = qa / "tracked-preflight"
            tracked_preflight = tracked_preflight_dir / "preflight.json"
            run(
                "tracked preflight",
                python(
                    final_scripts / "document_preflight.py",
                    "--source",
                    str(tracked_source),
                    "--candidate",
                    str(tracked_candidate),
                    "--final",
                    str(tracked_final),
                    "--output-dir",
                    str(tracked_preflight_dir),
                ),
            )
            tracked_render_dir = qa / "tracked-render"
            tracked_manifest = tracked_render_dir / "renderer.json"
            run(
                "tracked LibreOffice manifest render",
                python(
                    final_scripts / "render_docx_libreoffice.py",
                    "--input",
                    str(tracked_final),
                    "--output-dir",
                    str(tracked_render_dir),
                    "--manifest",
                    str(tracked_manifest),
                ),
            )
            run(
                "Word-authority requirement rejection",
                python(
                    final_scripts / "document_release.py",
                    "application-review",
                    "--preflight",
                    str(tracked_preflight),
                    "--renderer-manifest",
                    str(tracked_manifest),
                    "--output",
                    str(qa / "tracked-application-invalid.json"),
                    "--method",
                    "negative Word-authority regression",
                ),
                expect_success=False,
            )

        qa_ready = qa / "ready.json"
        run(
            "eight-gate QA review",
            python(
                final_scripts / "document_release.py",
                "qa-review",
                "--preflight",
                str(preflight),
                "--operation-completion",
                str(final_completion),
                "--handoff",
                str(handoff),
                "--render-evidence",
                str(render_evidence),
                "--scope-review",
                str(scope),
                "--structure-review",
                str(structure),
                "--native-policy-review",
                str(native),
                "--application-review",
                str(application),
                "--delivery-dir",
                str(delivery),
                "--deliverable",
                str(final),
                "--output",
                str(qa_ready),
            ),
        )
        release = qa / "release-record.json"
        run(
            "release",
            python(
                final_scripts / "document_release.py",
                "release",
                "--final",
                str(final),
                "--qa-review",
                str(qa_ready),
                "--output",
                str(release),
            ),
        )
        release_payload = json.loads(release.read_text(encoding="utf-8"))
        if release_payload.get("artifact_status") != "RELEASED":
            raise RuntimeError("release record does not contain RELEASED")
        published = root / "release-delivered.docx"
        delivery_record = qa / "delivery-record.json"
        input_at_start = root / "moved-original.docx"
        run(
            "same-directory delivery",
            python(
                final_scripts / "document_release.py",
                "deliver",
                "--release-record",
                str(release),
                "--input-at-start",
                str(input_at_start),
                "--destination",
                str(published),
                "--output",
                str(delivery_record),
            ),
        )
        delivery_payload = json.loads(delivery_record.read_text(encoding="utf-8"))
        if delivery_payload.get("status") != "DELIVERED":
            raise RuntimeError("delivery record does not contain DELIVERED")
        if published.parent != input_at_start.parent or sha256(published) != sha256(final):
            raise RuntimeError("same-directory delivery is not an exact copy")
        run(
            "live delivery verification",
            python(
                final_scripts / "document_release.py",
                "verify-delivery",
                "--delivery-record",
                str(delivery_record),
            ),
        )
        run(
            "delivery overwrite rejection",
            python(
                final_scripts / "document_release.py",
                "deliver",
                "--release-record",
                str(release),
                "--input-at-start",
                str(input_at_start),
                "--destination",
                str(published),
                "--output",
                str(qa / "delivery-overwrite.json"),
            ),
            expect_success=False,
        )
        run(
            "wrong-directory delivery rejection",
            python(
                final_scripts / "document_release.py",
                "deliver",
                "--release-record",
                str(release),
                "--input-at-start",
                str(input_at_start),
                "--destination",
                str(working / "wrong-dir.docx"),
                "--output",
                str(qa / "wrong-dir-delivery.json"),
            ),
            expect_success=False,
        )
        run(
            "delivery-record placement rejection",
            python(
                final_scripts / "document_release.py",
                "deliver",
                "--release-record",
                str(release),
                "--input-at-start",
                str(input_at_start),
                "--destination",
                str(root / "release-invalid-record-placement.docx"),
                "--output",
                str(delivery / "delivery-record-invalid.json"),
            ),
            expect_success=False,
        )
        published_bytes = published.read_bytes()
        published.write_bytes(published_bytes + b"\n")
        run(
            "tampered delivery rejection",
            python(
                final_scripts / "document_release.py",
                "verify-delivery",
                "--delivery-record",
                str(delivery_record),
            ),
            expect_success=False,
        )
        published.write_bytes(published_bytes)
        run(
            "release output overwrite rejection",
            python(
                final_scripts / "document_release.py",
                "release",
                "--final",
                str(final),
                "--qa-review",
                str(qa_ready),
                "--output",
                str(release),
            ),
            expect_success=False,
        )
        if sha256(source) != source_hash:
            raise RuntimeError("source hash changed during regression")
        completion_bytes = final_completion.read_bytes()
        final_completion.write_bytes(completion_bytes + b"\n")
        run(
            "tampered operation completion rejection",
            python(
                final_scripts / "document_release.py",
                "release",
                "--final",
                str(final),
                "--qa-review",
                str(qa_ready),
                "--output",
                str(qa / "release-after-operation-tamper.json"),
            ),
            expect_success=False,
        )
        final_completion.write_bytes(completion_bytes)
        comparison = preflight_dir / "comparison.json"
        with comparison.open("a", encoding="utf-8") as handle:
            handle.write("\n")
        run(
            "tampered evidence rejection",
            python(
                final_scripts / "document_release.py",
                "release",
                "--final",
                str(final),
                "--qa-review",
                str(qa_ready),
                "--output",
                str(qa / "release-after-tamper.json"),
            ),
            expect_success=False,
        )
        print(
            json.dumps(
                {
                    "status": "PASS",
                    "renderer": args.renderer,
                    "pages": len(pages),
                    "positive_release": True,
                    "same_directory_delivery": True,
                    "negative_checks": [
                        "operation_receipt_duplicate",
                        "fast_overwrite",
                        "missing_page",
                        "release_only_candidate_final_mismatch",
                        "visible_text_review_coverage",
                        "changed_part_coverage",
                        *(
                            ["word_authority_requirement"]
                            if args.renderer == "libreoffice"
                            else []
                        ),
                        "release_overwrite",
                        "delivery_overwrite",
                        "wrong_directory_delivery",
                        "delivery_record_placement",
                        "tampered_delivery",
                        "tampered_operation_completion",
                        "tampered_evidence",
                    ],
                    "source_sha256_preserved": source_hash,
                    "retained_workdir": str(root) if args.workdir else None,
                },
                ensure_ascii=False,
                indent=2,
            )
        )
        return 0
    except Exception as exc:
        print(json.dumps({"status": "FAIL", "error": str(exc), "workdir": str(root)}, indent=2))
        return 2
    finally:
        if temp_owner is not None:
            temp_owner.cleanup()


if __name__ == "__main__":
    raise SystemExit(main())
